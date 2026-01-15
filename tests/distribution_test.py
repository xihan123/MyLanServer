#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
在线填表功能测试脚本
支持多阶段批量提交、并发测试、全面错误测试
"""

import os
import sys
import json
import time
import random
import asyncio
import argparse
from datetime import datetime
from typing import Dict, List, Any, Optional, Tuple
from dataclasses import dataclass, field
from pathlib import Path
from io import BytesIO
import requests
import aiohttp
from faker import Faker

# 初始化 Faker
fake = Faker('zh_CN')

# 导入工具类模块
from utils.test_base import TestConfig, TestResult, TestCase, MockDatabase, BATCH_CONFIGS
from utils.http_client import HttpClient, AsyncHttpClient
from utils.report_generator import ReportGenerator


# ==================== 数据生成模块 ====================

class DataGenerator:
    """测试数据生成器"""

    def __init__(self, schema: Dict):
        self.schema = schema
        self.columns = schema.get('columns', [])

    def generate_field_value(self, field_def: Dict) -> Any:
        """根据字段定义生成测试数据"""
        field_name = field_def.get('name', '')
        field_type = field_def.get('type', 'Text')
        is_required = field_def.get('required', False)

        # 类型映射
        type_mapping = {
            'Text': self._generate_text,
            '数字': self._generate_number,
            'Number': self._generate_number,
            '日期': self._generate_date,
            'Date': self._generate_date,
            '布尔值': self._generate_boolean,
            'Boolean': self._generate_boolean,
            '双选框(是/否)': self._generate_boolean_text
        }

        generator = type_mapping.get(field_type, self._generate_text)
        value = generator(field_def)

        # 确保必填字段不为空
        if is_required and not value:
            value = self._generate_default_value(field_type)

        return value

    def _generate_text(self, field_def: Dict) -> str:
        """生成文本值"""
        field_name = field_def.get('name', '')

        # 根据字段名智能生成
        if '姓名' in field_name or 'name' in field_name.lower():
            return fake.name()
        elif '地址' in field_name or 'address' in field_name.lower():
            return fake.address()
        elif '邮箱' in field_name or 'email' in field_name.lower():
            return fake.email()
        elif '电话' in field_name or 'phone' in field_name.lower():
            return fake.phone_number()
        elif '部门' in field_name or 'department' in field_name.lower():
            return fake.company()
        else:
            return fake.sentence()

    def _generate_number(self, field_def: Dict) -> int:
        """生成数字值"""
        return random.randint(1, 100)

    def _generate_date(self, field_def: Dict) -> str:
        """生成日期值"""
        return fake.date_between(start_date='-30d', end_date='today').strftime('%Y-%m-%d')

    def _generate_boolean(self, field_def: Dict) -> bool:
        """生成布尔值"""
        return random.choice([True, False])

    def _generate_boolean_text(self, field_def: Dict) -> str:
        """生成双选框文本值（是/否）"""
        return random.choice(['true', 'false'])

    def _generate_default_value(self, field_type: str) -> Any:
        """生成默认值"""
        if field_type in ['数字', 'Number']:
            return 1
        elif field_type in ['日期', 'Date']:
            return '2026-01-10'
        elif field_type in ['布尔值', 'Boolean']:
            return True
        elif field_type == '双选框(是/否)':
            return 'true'
        else:
            return '测试'

    def generate_test_data(self) -> Dict:
        """生成完整的测试数据"""
        data = {}
        for column in self.columns:
            field_name = column.get('name', '')
            value = self.generate_field_value(column)
            data[field_name] = value
        return data

    def generate_multiple_test_data(self, count: int) -> List[Dict]:
        """生成多组测试数据"""
        return [self.generate_test_data() for _ in range(count)]


class TestFileGenerator:
    """测试文件生成器"""

    SUPPORTED_TYPES = ['pdf', 'docx', 'xlsx', 'txt', 'png', 'jpg', 'zip']

    @staticmethod
    def generate_txt(path: str, size: int = 1024) -> bytes:
        """生成文本文件"""
        content = '测试文本 ' * (size // 10 + 1)
        content = content.encode('utf-8')
        return content[:size]

    @staticmethod
    def generate_png(path: str, size: int = 1024) -> bytes:
        """生成 PNG 图片文件"""
        try:
            from PIL import Image
            import io

            # 计算图片尺寸（保持合理的宽高比）
            img_size = int((size * 8) ** 0.5)
            img_size = max(img_size, 100)  # 最小 100x100

            # 创建简单的图片
            img = Image.new('RGB', (img_size, img_size), color=(random.randint(0, 255), random.randint(0, 255), random.randint(0, 255)))

            # 保存到字节流
            buffer = io.BytesIO()
            img.save(buffer, format='PNG')
            content = buffer.getvalue()

            # 如果生成的文件太大，创建更小的图片
            if len(content) > size:
                img_size = int((size * 8) ** 0.5) // 2
                img = Image.new('RGB', (img_size, img_size), color=(random.randint(0, 255), random.randint(0, 255), random.randint(0, 255)))
                buffer = io.BytesIO()
                img.save(buffer, format='PNG')
                content = buffer.getvalue()

            return content
        except ImportError:
            # 如果 Pillow 不可用，生成简单的 PNG 文件头
            png_header = b'\x89PNG\r\n\x1a\n'
            return png_header + b'\x00' * (size - len(png_header))

    @staticmethod
    def generate_pdf(path: str, size: int = 1024) -> bytes:
        """生成 PDF 文件"""
        try:
            from reportlab.pdfgen import canvas
            from reportlab.lib.pagesizes import letter
            import io

            buffer = io.BytesIO()
            c = canvas.Canvas(buffer, pagesize=letter)

            # 添加一些文本
            c.drawString(100, 750, "测试 PDF 文件")
            c.drawString(100, 730, f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
            c.drawString(100, 710, f"测试内容: {'测试 ' * 100}")

            c.save()
            content = buffer.getvalue()

            # 如果生成的文件太大，调整内容
            if len(content) > size:
                buffer = io.BytesIO()
                c = canvas.Canvas(buffer, pagesize=letter)
                c.drawString(100, 750, "测试 PDF 文件")
                c.save()
                content = buffer.getvalue()

            return content
        except ImportError:
            # 如果 reportlab 不可用，生成简单的 PDF 文件头
            pdf_header = b'%PDF-1.4\n'
            return pdf_header + b'\x00' * (size - len(pdf_header))

    @staticmethod
    def generate_xlsx(path: str, size: int = 1024) -> bytes:
        """生成 Excel 文件"""
        try:
            from openpyxl import Workbook
            from openpyxl.writer.excel import save_virtual_workbook

            wb = Workbook()
            ws = wb.active
            ws.title = "测试"

            # 添加一些数据
            ws['A1'] = "测试 Excel 文件"
            ws['A2'] = f"生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"

            for i in range(3, 100):
                ws[f'A{i}'] = f"测试数据 {i}"

            content = save_virtual_workbook(wb)

            # 如果生成的文件太大，减少数据
            if len(content) > size:
                wb = Workbook()
                ws = wb.active
                ws['A1'] = "测试"
                content = save_virtual_workbook(wb)

            return content
        except ImportError:
            # 如果 openpyxl 不可用，生成简单的 XLSX 文件头
            xlsx_header = b'PK\x03\x04'
            return xlsx_header + b'\x00' * (size - len(xlsx_header))

    @staticmethod
    def generate_docx(path: str, size: int = 1024) -> bytes:
        """生成 Word 文件"""
        try:
            from docx import Document

            doc = Document()
            doc.add_heading('测试 Word 文件', 0)
            doc.add_paragraph(f'生成时间: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph('测试内容: ' + '测试 ' * 100)

            buffer = BytesIO()
            doc.save(buffer)
            content = buffer.getvalue()

            # 如果生成的文件太大，减少内容
            if len(content) > size:
                doc = Document()
                doc.add_paragraph('测试')
                buffer = BytesIO()
                doc.save(buffer)
                content = buffer.getvalue()

            return content
        except ImportError:
            # 如果 python-docx 不可用，生成简单的 DOCX 文件头
            docx_header = b'PK\x03\x04'
            return docx_header + b'\x00' * (size - len(docx_header))

    @staticmethod
    def generate_zip(path: str, size: int = 1024) -> bytes:
        """生成 ZIP 文件"""
        try:
            import zipfile

            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                # 添加一些文件
                zf.writestr('test.txt', '测试内容 ' * 100)

            content = buffer.getvalue()

            # 如果生成的文件太大，减少内容
            if len(content) > size:
                buffer = BytesIO()
                with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zf:
                    zf.writestr('test.txt', '测试')

                content = buffer.getvalue()

            return content
        except ImportError:
            # 如果 zipfile 不可用，生成简单的 ZIP 文件头
            zip_header = b'PK\x03\x04'
            return zip_header + b'\x00' * (size - len(zip_header))

    @classmethod
    def generate_file(cls, file_type: str, size: int = 1024) -> bytes:
        """生成指定类型的测试文件"""
        generators = {
            'txt': cls.generate_txt,
            'png': cls.generate_png,
            'jpg': cls.generate_png,  # 使用 PNG 生成器
            'pdf': cls.generate_pdf,
            'xlsx': cls.generate_xlsx,
            'docx': cls.generate_docx,
            'zip': cls.generate_zip
        }

        generator = generators.get(file_type.lower())
        if generator:
            return generator('', size)
        else:
            raise ValueError(f'不支持的文件类型: {file_type}')

    @classmethod
    def generate_all(cls, output_dir: str) -> Dict[str, str]:
        """生成所有类型的测试文件"""
        os.makedirs(output_dir, exist_ok=True)
        files = {}

        for file_type in cls.SUPPORTED_TYPES:
            filename = f'sample.{file_type}'
            filepath = os.path.join(output_dir, filename)

            # 生成文件内容
            content = cls.generate_file(file_type, size=random.randint(10 * 1024, 100 * 1024))

            # 保存文件
            with open(filepath, 'wb') as f:
                f.write(content)

            files[file_type] = filepath

        return files


# ==================== 测试用例 ====================

class SchemaTest(TestCase):
    """Schema 获取测试"""

    def __init__(self):
        super().__init__('Schema 获取测试', '测试获取表格结构定义接口')

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        success, data, duration = client.get_schema()

        if success:
            # 验证响应数据
            if 'title' in data and 'columns' in data:
                # 验证每个字段包含必需的属性
                valid_columns = True
                for column in data['columns']:
                    if 'name' not in column or 'type' not in column:
                        valid_columns = False
                        break

                if valid_columns:
                    return TestResult(
                        test_name=self.name,
                        passed=True,
                        message='Schema 获取成功',
                        duration=duration,
                        response_data=data,
                        details={
                            'title': data['title'],
                            'column_count': len(data['columns']),
                            'allow_attachment': data.get('allowAttachmentUpload', False)
                        }
                    )

        return TestResult(
            test_name=self.name,
            passed=False,
            message='Schema 获取失败',
            duration=duration,
            response_data=data if success else None,
            error=data.get('error', '未知错误') if not success else 'Schema 格式验证失败'
        )


class AttachmentListTest(TestCase):
    """附件列表获取测试"""

    def __init__(self):
        super().__init__('附件列表获取测试', '测试获取任务附件列表接口')

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        success, data, duration = client.get_attachments_list()

        if success:
            if 'attachments' in data:
                attachments = data['attachments']
                return TestResult(
                    test_name=self.name,
                    passed=True,
                    message=f'获取到 {len(attachments)} 个附件',
                    duration=duration,
                    response_data=data,
                    details={
                        'attachment_count': len(attachments),
                        'attachments': [
                            {
                                'id': att.get('id'),
                                'file_name': att.get('fileName'),
                                'display_name': att.get('displayName'),
                                'file_size': att.get('fileSize')
                            }
                            for att in attachments
                        ]
                    }
                )

        return TestResult(
            test_name=self.name,
            passed=False,
            message='附件列表获取失败',
            duration=duration,
            response_data=data if success else None,
            error=data.get('error', '未知错误') if not success else '响应格式验证失败'
        )


class AttachmentDownloadTest(TestCase):
    """附件下载测试"""

    def __init__(self):
        super().__init__('附件下载测试', '测试下载任务附件接口')

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        # 先获取附件列表
        success, data, _ = client.get_attachments_list()

        if not success or 'attachments' not in data:
            return TestResult(
                test_name=self.name,
                passed=False,
                message='无法获取附件列表',
                duration=time.time() - start_time,
                error=data.get('error', '未知错误')
            )

        attachments = data['attachments']

        if not attachments:
            return TestResult(
                test_name=self.name,
                passed=True,
                message='任务没有附件，跳过下载测试',
                duration=time.time() - start_time,
                details={'attachment_count': 0}
            )

        # 下载所有附件
        download_results = []
        total_duration = 0

        for attachment in attachments:
            attachment_id = attachment.get('id')
            file_name = attachment.get('fileName', f'attachment_{attachment_id}')

            success, content, duration = client.download_attachment(attachment_id)
            total_duration += duration

            if success and len(content) > 0:
                # 保存下载的文件
                save_path = os.path.join(config.downloads_dir, file_name)
                with open(save_path, 'wb') as f:
                    f.write(content)

                download_results.append({
                    'id': attachment_id,
                    'file_name': file_name,
                    'size': len(content),
                    'status': 'success'
                })
            else:
                download_results.append({
                    'id': attachment_id,
                    'file_name': file_name,
                    'size': 0,
                    'status': 'failed'
                })

        # 检查是否所有附件都下载成功
        all_success = all(result['status'] == 'success' for result in download_results)

        return TestResult(
            test_name=self.name,
            passed=all_success,
            message=f'下载了 {len(download_results)} 个附件',
            duration=total_duration,
            details={
                'total_attachments': len(attachments),
                'successful_downloads': sum(1 for r in download_results if r['status'] == 'success'),
                'failed_downloads': sum(1 for r in download_results if r['status'] == 'failed'),
                'downloads': download_results
            }
        )


class DataGenerationTest(TestCase):
    """数据生成测试"""

    def __init__(self, schema: Dict):
        super().__init__('数据生成测试', '测试根据 Schema 生成测试数据')
        self.schema = schema

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        try:
            generator = DataGenerator(self.schema)
            test_data = generator.generate_test_data()

            # 验证生成的数据
            validation_results = []
            for column in self.schema.get('columns', []):
                field_name = column.get('name', '')
                field_type = column.get('type', 'Text')
                is_required = column.get('required', False)

                if field_name not in test_data:
                    validation_results.append({
                        'field': field_name,
                        'status': 'failed',
                        'reason': '字段不存在'
                    })
                    continue

                value = test_data[field_name]

                if is_required and not value:
                    validation_results.append({
                        'field': field_name,
                        'status': 'failed',
                        'reason': '必填字段为空'
                    })
                    continue

                validation_results.append({
                    'field': field_name,
                    'status': 'success',
                    'type': field_type,
                    'value': value
                })

            all_valid = all(result['status'] == 'success' for result in validation_results)

            return TestResult(
                test_name=self.name,
                passed=all_valid,
                message='数据生成成功',
                duration=time.time() - start_time,
                details={
                    'generated_data': test_data,
                    'validation': validation_results
                }
            )
        except Exception as e:
            return TestResult(
                test_name=self.name,
                passed=False,
                message='数据生成失败',
                duration=time.time() - start_time,
                error=str(e)
            )


class FormSubmissionTest(TestCase):
    """表单提交测试"""

    def __init__(self, name: str, description: str, include_attachments: bool = False, attachment_count: int = 0):
        super().__init__(name, description)
        self.include_attachments = include_attachments
        self.attachment_count = attachment_count

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        try:
            # 生成测试数据
            success, schema_data, _ = client.get_schema()
            if not success:
                return TestResult(
                    test_name=self.name,
                    passed=False,
                    message='无法获取 Schema',
                    duration=time.time() - start_time,
                    error=schema_data.get('error', '未知错误')
                )

            generator = DataGenerator(schema_data)
            form_data = generator.generate_test_data()

            # 准备提交数据
            submit_data = {
                'name': config.test_user['name'],
                'contact': config.test_user['contact'],
                'department': config.test_user['department'],
                'jsonData': form_data
            }

            # 准备附件
            files = None
            if self.include_attachments and self.attachment_count > 0:
                files = []
                for i in range(self.attachment_count):
                    # 只使用任务允许的格式：.jpg, .png（不使用 jpeg，因为生成器会创建 PNG 内容）
                    file_type = random.choice(['png', 'jpg'])
                    content = TestFileGenerator.generate_file(file_type, size=random.randint(10 * 1024, 50 * 1024))
                    filename = f'test_{i + 1}.{file_type}'
                    files.append(('attachment', filename, content))

            # 提交表单
            success, response_data, duration = client.submit_form(submit_data, files)

            if success:
                # 验证响应
                if 'message' in response_data and 'filename' in response_data:
                    # 保存到模拟数据库
                    db.insert('test_submissions', {
                        'name': submit_data['name'],
                        'contact': submit_data['contact'],
                        'department': submit_data['department'],
                        'data': form_data,
                        'attachment_count': self.attachment_count,
                        'status': 'success',
                        'response': response_data
                    })

                    return TestResult(
                        test_name=self.name,
                        passed=True,
                        message='表单提交成功',
                        duration=duration,
                        response_data=response_data,
                        details={
                            'submitter': response_data.get('submitter'),
                            'contact': response_data.get('contact'),
                            'department': response_data.get('department'),
                            'filename': response_data.get('filename'),
                            'attachment_count': response_data.get('attachmentCount', 0)
                        }
                    )

            return TestResult(
                test_name=self.name,
                passed=False,
                message='表单提交失败',
                duration=duration,
                response_data=response_data if success else None,
                error=response_data.get('error', '未知错误') if not success else '响应格式验证失败'
            )
        except Exception as e:
            return TestResult(
                test_name=self.name,
                passed=False,
                message='表单提交失败',
                duration=time.time() - start_time,
                error=str(e)
            )


class BatchSubmissionTest(TestCase):
    """批量提交测试"""

    def __init__(self, count: int, concurrent: int = 1):
        super().__init__(
            f'批量提交测试 ({count}次提交, {concurrent}并发)',
            f'测试批量提交表单功能'
        )
        self.count = count
        self.concurrent = concurrent

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        try:
            # 获取 Schema
            success, schema_data, _ = client.get_schema()
            if not success:
                return TestResult(
                    test_name=self.name,
                    passed=False,
                    message='无法获取 Schema',
                    duration=time.time() - start_time,
                    error=schema_data.get('error', '未知错误')
                )

            generator = DataGenerator(schema_data)

            if self.concurrent > 1:
                # 并发测试
                loop = asyncio.get_event_loop()
                result = loop.run_until_complete(self._run_concurrent_test(config, generator, db))
                return result
            else:
                # 单线程测试
                return self._run_sequential_test(client, config, generator, db)

        except Exception as e:
            return TestResult(
                test_name=self.name,
                passed=False,
                message='批量提交测试失败',
                duration=time.time() - start_time,
                error=str(e)
            )

    def _run_sequential_test(self, client: HttpClient, config: TestConfig, generator: DataGenerator, db: MockDatabase) -> TestResult:
        """运行单线程批量测试"""
        start_time = time.time()
        results = []

        for i in range(self.count):
            # 生成测试数据
            form_data = generator.generate_test_data()

            # 准备提交数据
            submit_data = {
                'name': f'测试用户{i + 1}',
                'contact': str(random.randint(10000000000, 99999999999)),
                'department': random.choice(['技术部', '人事部', '财务部', '市场部']),
                'jsonData': form_data
            }

            # 提交表单
            success, response_data, duration = client.submit_form(submit_data)

            results.append({
                'index': i + 1,
                'name': submit_data['name'],
                'success': success,
                'duration': duration,
                'error': response_data.get('error') if not success else None
            })

            # 保存到模拟数据库
            if success:
                db.insert('test_submissions', {
                    'name': submit_data['name'],
                    'contact': submit_data['contact'],
                    'department': submit_data['department'],
                    'data': form_data,
                    'attachment_count': 0,
                    'status': 'success',
                    'response': response_data
                })

        total_duration = time.time() - start_time
        success_count = sum(1 for r in results if r['success'])

        return TestResult(
            test_name=self.name,
            passed=success_count > 0,
            message=f'成功提交 {success_count}/{self.count} 次',
            duration=total_duration,
            details={
                'total_submissions': self.count,
                'successful_submissions': success_count,
                'failed_submissions': self.count - success_count,
                'average_duration': sum(r['duration'] for r in results) / len(results) if results else 0,
                'max_duration': max(r['duration'] for r in results) if results else 0,
                'min_duration': min(r['duration'] for r in results) if results else 0,
                'results': results[:10]  # 只保留前10条详细结果
            }
        )

    async def _run_concurrent_test(self, config: TestConfig, generator: DataGenerator, db: MockDatabase) -> TestResult:
        """运行并发批量测试"""
        start_time = time.time()

        # 创建异步客户端
        async_client = AsyncHttpClient(config)

        # 创建异步任务
        tasks = []
        for i in range(self.count):
            # 生成测试数据
            form_data = generator.generate_test_data()

            # 准备提交数据
            submit_data = {
                'name': f'测试用户{i + 1}',
                'contact': str(random.randint(10000000000, 99999999999)),
                'department': random.choice(['技术部', '人事部', '财务部', '市场部']),
                'jsonData': form_data
            }

            task = async_client.submit_form_async(submit_data)
            tasks.append((i + 1, submit_data['name'], task))

        # 执行并发任务
        results = []
        completed_tasks = await asyncio.gather(*[task for _, _, task in tasks], return_exceptions=True)

        for (index, name, _), result in zip(tasks, completed_tasks):
            if isinstance(result, Exception):
                results.append({
                    'index': index,
                    'name': name,
                    'success': False,
                    'duration': 0,
                    'error': str(result)
                })
            else:
                success, response_data, duration = result
                results.append({
                    'index': index,
                    'name': name,
                    'success': success,
                    'duration': duration,
                    'error': response_data.get('error') if not success else None
                })

                # 保存到模拟数据库
                if success:
                    db.insert('test_submissions', {
                        'name': name,
                        'contact': '',
                        'department': '',
                        'data': {},
                        'attachment_count': 0,
                        'status': 'success',
                        'response': response_data
                    })

        total_duration = time.time() - start_time
        success_count = sum(1 for r in results if r['success'])

        return TestResult(
            test_name=self.name,
            passed=success_count > 0,
            message=f'成功提交 {success_count}/{self.count} 次',
            duration=total_duration,
            details={
                'total_submissions': self.count,
                'successful_submissions': success_count,
                'failed_submissions': self.count - success_count,
                'concurrent': self.concurrent,
                'average_duration': sum(r['duration'] for r in results) / len(results) if results else 0,
                'max_duration': max(r['duration'] for r in results) if results else 0,
                'min_duration': min(r['duration'] for r in results) if results else 0,
                'results': results[:10]  # 只保留前10条详细结果
            }
        )


class ErrorHandlingTest(TestCase):
    """错误处理测试"""

    def __init__(self):
        super().__init__('错误处理测试', '测试各种错误场景的处理')

    def execute(self, client: HttpClient, config: TestConfig, db: MockDatabase) -> TestResult:
        start_time = time.time()

        # 获取 Schema
        success, schema_data, _ = client.get_schema()
        if not success:
            return TestResult(
                test_name=self.name,
                passed=False,
                message='无法获取 Schema',
                duration=time.time() - start_time,
                error=schema_data.get('error', '未知错误')
            )

        generator = DataGenerator(schema_data)
        form_data = generator.generate_test_data()

        # 错误测试场景
        error_scenarios = [
            {'name': '无效密码', 'password': 'wrong_password', 'expected_status': 'fail'},
            {'name': '无效 Slug', 'slug': 'INVALID_SLUG', 'expected_status': 'fail'},
            {'name': '缺少姓名', 'name': None, 'expected_status': 'fail'},
            {'name': '缺少联系方式', 'contact': None, 'expected_status': 'fail'},
            {'name': '联系方式过短', 'contact': '12', 'expected_status': 'fail'},
            {'name': '联系方式过长', 'contact': '123456789012345', 'expected_status': 'fail'},
            {'name': '缺少表单数据', 'jsonData': None, 'expected_status': 'fail'},
            {'name': '必填字段缺失', 'jsonData': {}, 'expected_status': 'fail'},
        ]

        test_results = []

        for scenario in error_scenarios:
            try:
                # 准备测试数据
                submit_data = {
                    'name': scenario.get('name', config.test_user['name']) if scenario.get('name') is not None else None,
                    'contact': scenario.get('contact', config.test_user['contact']),
                    'department': config.test_user['department'],
                    'jsonData': scenario.get('jsonData', form_data)
                }

                # 使用自定义密码
                password = scenario.get('password', config.password)
                submit_data['password'] = password

                # 使用自定义 Slug
                slug = scenario.get('slug', config.slug)

                # 临时修改配置
                original_slug = config.slug
                config.slug = slug

                # 提交表单
                success, response_data, duration = client.submit_form(submit_data)

                # 恢复配置
                config.slug = original_slug

                # 判断测试是否通过（预期失败且实际失败）
                expected_fail = scenario['expected_status'] == 'fail'
                actual_fail = not success

                test_passed = (expected_fail and actual_fail)

                test_results.append({
                    'scenario': scenario['name'],
                    'passed': test_passed,
                    'expected': scenario['expected_status'],
                    'actual': 'fail' if not success else 'success',
                    'error': response_data.get('error') if not success else None
                })

                # 保存到模拟数据库
                db.insert('test_errors', {
                    'scenario': scenario['name'],
                    'passed': test_passed,
                    'error': response_data.get('error') if not success else None,
                    'status': 'error_test'
                })

            except Exception as e:
                test_results.append({
                    'scenario': scenario['name'],
                    'passed': False,
                    'expected': scenario['expected_status'],
                    'actual': 'exception',
                    'error': str(e)
                })

        total_duration = time.time() - start_time
        passed_count = sum(1 for r in test_results if r['passed'])

        return TestResult(
            test_name=self.name,
            passed=passed_count > 0,
            message=f'错误测试通过 {passed_count}/{len(test_results)}',
            duration=total_duration,
            details={
                'total_scenarios': len(error_scenarios),
                'passed_scenarios': passed_count,
                'failed_scenarios': len(error_scenarios) - passed_count,
                'results': test_results
            }
        )


# ==================== 测试执行引擎 ====================

class TestRunner:
    """测试执行引擎"""

    def __init__(self, config: TestConfig):
        self.config = config
        self.client = HttpClient(config)
        self.db = MockDatabase()
        self.results: List[TestResult] = []

    def register_test(self, test_case: TestCase):
        """注册测试用例"""
        test_case.setup()
        result = test_case.execute(self.client, self.config, self.db)
        test_case.teardown()
        self.results.append(result)
        return result

    def run_all(self) -> List[TestResult]:
        """运行所有测试"""
        print(f'\n开始测试...')
        print(f'Base API: {self.config.base_api}')
        print(f'Slug: {self.config.slug}')
        print(f'批量提交: {self.config.batch_count} 次, {self.config.concurrent} 并发')
        print(f'-' * 60)

        # 1. Schema 获取测试
        result = self.register_test(SchemaTest())
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 2. 附件列表获取测试
        result = self.register_test(AttachmentListTest())
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 3. 附件下载测试
        result = self.register_test(AttachmentDownloadTest())
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 4. 数据生成测试
        success, schema_data, _ = self.client.get_schema()
        if success:
            result = self.register_test(DataGenerationTest(schema_data))
            print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 5. 表单提交测试（无附件）
        result = self.register_test(FormSubmissionTest('表单提交测试（无附件）', '测试不带附件的表单提交', False, 0))
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 6. 表单提交测试（单个附件）
        result = self.register_test(FormSubmissionTest('表单提交测试（单个附件）', '测试带单个附件的表单提交', True, 1))
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 7. 表单提交测试（多个附件）
        result = self.register_test(FormSubmissionTest('表单提交测试（多个附件）', '测试带多个附件的表单提交', True, 3))
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 8. 批量提交测试
        result = self.register_test(BatchSubmissionTest(self.config.batch_count, self.config.concurrent))
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        # 9. 错误处理测试
        result = self.register_test(ErrorHandlingTest())
        print(f'✓ {result.test_name}: {"通过" if result.passed else "失败"} ({result.duration:.2f}s)')

        print(f'-' * 60)
        print(f'测试完成!')

        return self.results


# ==================== 主程序入口 ====================

def main():
    """主程序入口"""
    parser = argparse.ArgumentParser(description='在线填表功能测试脚本')
    parser.add_argument('--base-api',default='http://192.168.0.100:8080', help='Base API URL (例如: http://192.168.0.100:8080)')
    parser.add_argument('--slug', required=True, help='任务 Slug')
    parser.add_argument('--password', required=True, help='访问密码')
    parser.add_argument('--name', default='测试用户', help='测试用户姓名')
    parser.add_argument('--contact', default='12345678901', help='测试用户联系方式')
    parser.add_argument('--department', default='技术部', help='测试用户部门')
    parser.add_argument('--batch', choices=['single', 'small', 'medium', 'large'], default='small', help='批量配置预设')
    parser.add_argument('--count', type=int, help='自定义提交次数')
    parser.add_argument('--concurrent', type=int, help='自定义并发数')
    parser.add_argument('--output-dir', default='test_reports', help='报告输出目录')
    parser.add_argument('--no-files', action='store_true', help='不自动生成测试文件')

    args = parser.parse_args()

    # 确定批量配置
    if args.count:
        batch_count = args.count
    else:
        batch_count = BATCH_CONFIGS[args.batch]['count']

    if args.concurrent:
        concurrent = args.concurrent
    else:
        concurrent = BATCH_CONFIGS[args.batch]['concurrent']

    # 创建配置
    config = TestConfig(
        base_api=args.base_api,
        slug=args.slug,
        password=args.password,
        test_user={
            'name': args.name,
            'contact': args.contact,
            'department': args.department
        },
        batch_count=batch_count,
        concurrent=concurrent,
        output_dir=args.output_dir
    )

    # 自动生成测试文件
    if not args.no_files:
        print('\n正在生成测试文件...')
        TestFileGenerator.generate_all(config.test_files_dir)
        print(f'✓ 测试文件已生成到 {config.test_files_dir}')

    # 运行测试
    runner = TestRunner(config)
    results = runner.run_all()

    # 生成报告
    print('\n正在生成测试报告...')

    # 获取 Schema 信息作为额外信息
    success, schema_data, _ = HttpClient(config).get_schema()
    additional_info = {'schema': schema_data if success else None}

    # 生成报告
    report_generator = ReportGenerator(config.output_dir)
    report_content = report_generator.generate_markdown_report(results, config, additional_info)

    # 保存报告
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    report_filename = f'test_report_{timestamp}.md'
    report_path = report_generator.save_report(report_content, report_filename)

    print(f'✓ 测试报告已保存到 {report_path}')

    # 打印统计信息
    print(f'\n{"=" * 60}')
    print(f'测试统计:')
    print(f'{"=" * 60}')
    print(f'总测试用例: {len(results)}')
    print(f'通过: {sum(1 for r in results if r.passed)}')
    print(f'失败: {sum(1 for r in results if not r.passed)}')
    print(f'通过率: {sum(1 for r in results if r.passed) / len(results) * 100:.1f}%')
    print(f'{"=" * 60}\n')

    # 返回退出码
    return 0 if all(r.passed for r in results) else 1


if __name__ == '__main__':
    sys.exit(main())